import json
import os
from pathlib import Path
import time
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document
from resume_score import final_score

load_dotenv()

api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    raise ValueError("api error")
client = Groq(api_key=api_key)
model="llama-3.3-70b-versatile"


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None

